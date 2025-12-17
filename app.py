import streamlit as st
import pandas as pd
import numpy as np
import io
import os
from zipfile import ZipFile
from docx import Document
import csv

# -----------------------------------------------------------
# Paths (relative to the app folder)
# -----------------------------------------------------------
STATE_CARRIER_PATH = "HD_carrier_guide.xlsx"
TEMPLATE_PATH = "BOL_template.docx"
OUTPUT_FOLDER = "BOL_created"

# Create output folder if it doesn't exist
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# -----------------------------------------------------------
# Load preloaded carrier guide
# -----------------------------------------------------------
@st.cache_data
def load_state_carrier():
    df = pd.read_excel(STATE_CARRIER_PATH)
    df.columns = df.columns.str.strip().str.replace(r"\s+", "", regex=True)
    TN_carrier_df = df[df['OriginState'] == 'TN'].reset_index(drop=True)
    columns_to_keep = ['SupplierIBtoDC/StoreCarrier', 'ResidentialDeliveryCarrier(Hd.com)', 'DestinationState']
    TN_carrier_clean = TN_carrier_df[columns_to_keep].copy()

    mapping = {
        'AACT': 'AAA Cooper Transportation',
        'EXLA': 'Estes Express Lines',
        'CTII': 'Central Transport Inc.',
        'ABFS': 'ABF',
        'RNLO': 'R&L Carriers'
    }
    TN_carrier_clean['ShippingCodeStores'] = TN_carrier_clean['SupplierIBtoDC/StoreCarrier'].map(mapping)
    TN_carrier_clean['ShippingCodeHomeDelivery'] = TN_carrier_clean['ResidentialDeliveryCarrier(Hd.com)'].map(mapping)
    
    return TN_carrier_clean

# -----------------------------------------------------------
# Word template filling function
# -----------------------------------------------------------
def fill_template(template_path, output_path, replacements):
    doc = Document(template_path)
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))
    doc.save(output_path)

# -----------------------------------------------------------
# Streamlit UI
# -----------------------------------------------------------
st.set_page_config(page_title="📄 BOL Generator", layout="centered")
st.title("📄 BOL Generation Tool")
st.write("Upload SAP LTL Cleaned file and CommerceHub CSV to generate BOL Word documents.")

st.markdown("---")

# Upload files
uploaded_ltl = st.file_uploader("📤 Upload SAP LTL Cleaned Excel file", type=["xlsx"])
uploaded_csv = st.file_uploader("📤 Upload CommerceHub CSV file", type=["csv"])

TN_carrier_clean = load_state_carrier()

if uploaded_ltl and uploaded_csv:
    if st.button("▶️ Process & Generate BOLs"):
        try:
            # Load SAP LTL Cleaned
            df_LTL_grouped = pd.read_excel(uploaded_ltl)
            df_LTL_grouped['Purchase order no.'] = df_LTL_grouped['Purchase order no.'].astype(int)

            # Load CommerceHub CSV
            uploaded_csv.seek(0) # Ensure file pointer is at the start

            # Read CSV, skipping first 3 rows
            df_csv = pd.read_csv(uploaded_csv, skiprows=3, encoding='utf-8', engine='python')
            
            df_csv.columns = df_csv.columns.str.strip().str.replace(r"\s+", "", regex=True)
            df_csv['HD_Store'] = df_csv['ShipToAddress1'].str.extract(r'Store #(\d{3,})')
            df_csv['ShipToAddress'] = np.where(
                df_csv['ShipToAddress1'].str.contains('THD', na=False),
                df_csv['ShipToAddress2'],
                df_csv['ShipToAddress1']
            )

            # Merge carriers
            df_chub = pd.merge(
                df_csv,
                TN_carrier_clean,
                left_on='ShipToState',
                right_on='DestinationState',
                how='left'
            )
            df_chub.drop(columns='DestinationState', inplace=True)
            df_chub['SCAC'] = np.where(
                df_chub['HD_Store'].notna(),
                df_chub['SupplierIBtoDC/StoreCarrier'],
                df_chub['ResidentialDeliveryCarrier(Hd.com)']
            )
            df_chub['Carrier_name'] = np.where(
                df_chub['HD_Store'].notna(),
                df_chub['ShippingCodeStores'],
                df_chub['ShippingCodeHomeDelivery']
            )

            # Merge with LTL grouped
            df_BOL = pd.merge(
                df_LTL_grouped,
                df_chub,
                left_on=['Purchase order no.'],
                right_on=['PONumber'],
                how='inner'
            )

            # Generate Word files with progress bar
            created_files = []
            progress_bar = st.progress(0)
            total_rows = len(df_BOL)

            for i, row in df_BOL.iterrows():
                dn = row.get("DN","") #get DN for filename
                dn = str(dn).strip().replace("/","_").replace("//"."_") #make sure filename is clean
                                                              
                replacements = {
                    "{{CARRIER NAME}}": row.get("Carrier_name", ""),
                    "{{CUSTOMER NAME}}": row.get("ShipToName", ""),
                    "{{HD_STORE}}": row.get("ShipToAddress1", "") if pd.notna(row.get("HD_Store")) else "",
                    "{{ADRESS}}": row.get("ShipToAddress", ""),
                    "{{CITY}}": row.get("ShipToCity", ""),
                    "{{STATE}}": row.get("ShipToState", ""),
                    "{{ZIP CODE}}": row.get("ShipToPostalCode", ""),
                    "{{PHONE NUMBER}}": row.get("ShipToDayPhone", ""),
                    "{{SCAC}}": row.get("SCAC", ""),
                    "{{PO_NUMBER}}": row.get("PONumber", ""),
                    "{{NUM_PACKAGES}}": str(row.get("Order Quantity", "")),
                    "{{WEIGHT}}": str(row.get("Gross weight", "")),
                    "{{CUSTOMER ORDER}}": row.get("CustomerOrderNumber", ""),
                    "{{DELIVERY NUMBER}}": row.get("DN", ""),
                    "{{QTY_1}}": str(row.get("Pallet_qty", "")),
                    "{{QTY_PACK}}": str(row.get("Order Quantity", ""))
                }

                output_file = os.path.join(OUTPUT_FOLDER, f"BOL_{dn}.docx")
                fill_template(TEMPLATE_PATH, output_file, replacements)
                created_files.append(output_file)

                # Update progress bar
                progress_bar.progress((i + 1) / total_rows)

            # Create ZIP
            zip_buffer = io.BytesIO()
            with ZipFile(zip_buffer, "w") as zipf:
                for file in created_files:
                    zipf.write(file, os.path.basename(file))
            zip_buffer.seek(0)

            st.success(f"✅ Created {len(created_files)} BOLs!")
            st.download_button(
                "⬇️ Download All BOLs (ZIP)",
                data=zip_buffer,
                file_name="BOLs.zip",
                mime="application/zip"
            )

        except Exception as e:
            st.error(f"❌ Error processing files: {e}")
else:
    st.info("Upload both SAP LTL Cleaned Excel and CommerceHub CSV files to start processing.")

